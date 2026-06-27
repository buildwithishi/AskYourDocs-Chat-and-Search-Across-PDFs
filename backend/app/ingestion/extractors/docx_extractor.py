import io

from app.ingestion.extractors.base import ExtractedDocument, ExtractedPage, TextExtractor


class DocxExtractor(TextExtractor):
    def extract(self, content: bytes) -> ExtractedDocument:
        from docx import Document as DocxDocument  # deferred: heavy import, loaded lazily

        docx_document = DocxDocument(io.BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in docx_document.paragraphs)
        return ExtractedDocument(pages=[ExtractedPage(page_number=None, text=text)], page_count=None)
